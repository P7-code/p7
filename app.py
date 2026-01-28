"""
安天投标文件智能分析系统 - Web界面
"""
import os
import sys
import json
import tempfile
from typing import Dict, Any
from datetime import datetime
import streamlit as st

# 添加src到Python路径
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from graphs.graph import main_graph
from utils.file.file import File
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# 页面配置
st.set_page_config(
    page_title="安天投标文件智能分析系统",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 自定义CSS
st.markdown("""
<style>
    .main-title {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.8rem;
        color: #2e7d32;
        margin-top: 1.5rem;
        margin-bottom: 1rem;
    }
    .info-box {
        background-color: #e3f2fd;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #c8e6c9;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)


def save_uploaded_file(uploaded_file) -> str:
    """保存上传的文件到临时目录"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{uploaded_file.name}") as tmp_file:
            tmp_file.write(uploaded_file.getbuffer())
            return tmp_file.name
    except Exception as e:
        st.error(f"文件保存失败: {str(e)}")
        return None


def display_checklist_result(checklist: Dict[str, Any], section_title: str, color_class: str = "info-box"):
    """显示检查清单结果"""
    st.markdown(f"### {section_title}")
    st.markdown(f'<div class="{color_class}">', unsafe_allow_html=True)
    
    # 如果是字符串，直接显示
    if isinstance(checklist, str):
        st.markdown(checklist)
    elif isinstance(checklist, dict):
        for key, value in checklist.items():
            if isinstance(value, list):
                st.markdown(f"**{key}:**")
                for item in value:
                    if isinstance(item, dict):
                        for k, v in item.items():
                            st.markdown(f"- {k}: {v}")
                    else:
                        st.markdown(f"- {item}")
            else:
                st.markdown(f"**{key}:** {value}")
    elif isinstance(checklist, list):
        for item in checklist:
            if isinstance(item, dict):
                for k, v in item.items():
                    st.markdown(f"**{k}:** {v}")
            else:
                st.markdown(f"- {item}")
    
    st.markdown("</div>", unsafe_allow_html=True)


def generate_docx_report(result: Dict[str, Any]) -> bytes:
    """
    生成docx格式的分析报告

    Args:
        result: 分析结果字典

    Returns:
        docx文件的字节数据
    """
    doc = Document()

    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题
    title = doc.add_heading('投标文件智能分析报告', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 生成时间
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()

    # 废标项检查
    doc.add_heading('一、废标项检查', level=1)
    invalid_items = result.get("invalid_items_check", "")
    if invalid_items:
        add_content_to_docx(doc, invalid_items)

    # 商务得分检查
    doc.add_heading('二、商务得分检查', level=1)
    commercial_score = result.get("commercial_score_check", "")
    if commercial_score:
        add_content_to_docx(doc, commercial_score)

    # 技术方案检查
    doc.add_heading('三、技术方案检查', level=1)
    technical_plan = result.get("technical_plan_check", "")
    if technical_plan:
        add_content_to_docx(doc, technical_plan)

    # 指标应答检查
    doc.add_heading('四、指标应答检查', level=1)
    indicator_response = result.get("indicator_response_check", "")
    if indicator_response:
        add_content_to_docx(doc, indicator_response)

    # 技术得分检查
    doc.add_heading('五、技术得分检查', level=1)
    technical_score = result.get("technical_score_check", "")
    if technical_score:
        add_content_to_docx(doc, technical_score)

    # 文件结构检查
    doc.add_heading('六、文件结构检查', level=1)
    bid_structure = result.get("bid_structure_check", "")
    if bid_structure:
        add_content_to_docx(doc, bid_structure)

    # 修改建议汇总
    doc.add_heading('七、修改建议汇总', level=1)
    summary = result.get("final_modification_suggestions", "")
    if summary:
        add_content_to_docx(doc, summary)

    # 保存到字节流
    from io import BytesIO
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    return doc_stream.getvalue()


def add_content_to_docx(doc: Document, content: str):
    """
    将内容添加到docx文档中

    Args:
        doc: docx文档对象
        content: 要添加的内容
    """
    lines = content.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 判断是否为标题（以===或##开头）
        if line.startswith('===') or line.startswith('#'):
            level = 2
            if line.startswith('===') and line.count('=') > 5:
                level = 1
            elif line.startswith('###'):
                level = 3
            doc.add_heading(line.lstrip('= #'), level=level)
        # 判断是否为列表项（以数字或-开头）
        elif line[0].isdigit() or (line[0] == '-' and len(line) > 1 and line[1].isspace()):
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)


def main():
    """主函数"""
    # 标题
    st.markdown('<h1 class="main-title">📊 安天投标文件智能分析系统</h1>', unsafe_allow_html=True)
    
    # 侧边栏说明
    with st.sidebar:
        st.markdown("## 📖 使用说明")
        st.markdown("""
        1. 上传招标文件（PDF/Word）
        2. 上传投标文件（PDF/Word/PPT）
        3. 点击"开始分析"按钮
        4. 查看分析结果和修改建议
        """)
        
        st.markdown("---")
        st.markdown("## 💡 系统功能")
        st.markdown("""
        - ✅ 废标项检测
        - ✅ 商务得分检查
        - ✅ 技术方案评估
        - ✅ 指标应答验证
        - ✅ 技术得分点分析
        - ✅ 文件结构检查
        - ✅ 生成修改建议
        """)
        
        st.markdown("---")
        st.markdown("## ⚠️ 注意事项")
        st.markdown("""
        - 仅支持PDF、Word、PPT格式
        - 文件大小不超过100MB
        - 分析过程可能需要几分钟
        """)
    
    # 主内容区
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown('<div class="section-header">📄 招标文件</div>', unsafe_allow_html=True)
        tender_file = st.file_uploader(
            "上传招标文件",
            type=['pdf', 'docx', 'doc', 'pptx', 'ppt'],
            key="tender_file"
        )
        if tender_file:
            st.success(f"已选择: {tender_file.name}")
    
    with col2:
        st.markdown('<div class="section-header">📝 投标文件</div>', unsafe_allow_html=True)
        bid_file = st.file_uploader(
            "上传投标文件",
            type=['pdf', 'docx', 'doc', 'pptx', 'ppt'],
            key="bid_file"
        )
        if bid_file:
            st.success(f"已选择: {bid_file.name}")
    
    # 分析按钮
    st.markdown("---")
    analyze_button = st.button("🚀 开始分析", type="primary", use_container_width=True)
    
    if analyze_button:
        if not tender_file or not bid_file:
            st.error("❌ 请先上传招标文件和投标文件！")
            return
        
        # 保存文件
        with st.spinner("正在保存文件..."):
            tender_path = save_uploaded_file(tender_file)
            bid_path = save_uploaded_file(bid_file)
            
            if not tender_path or not bid_path:
                st.error("文件保存失败！")
                return
        
        # 准备输入
        try:
            input_data = {
                "tender_file": {
                    "url": tender_path,
                    "file_type": "document"
                },
                "bid_file": {
                    "url": bid_path,
                    "file_type": "document"
                }
            }
            
            st.success("文件准备就绪，开始分析...")
            
            # 运行工作流
            with st.spinner("正在进行六维分析，请稍候..."):
                result = main_graph.invoke(input_data)
            
            # 显示结果
            st.markdown('<h2 class="section-header">📋 分析结果</h2>', unsafe_allow_html=True)

            # 废标项检测结果
            if result.get("invalid_items_check"):
                invalid_items = result["invalid_items_check"]
                # 检查是否包含废标风险关键词
                if "未发现废标项" in invalid_items or "无废标风险" in invalid_items or "恭喜" in invalid_items:
                    st.markdown('<div class="success-box">✅ 未发现废标项，恭喜！</div>', unsafe_allow_html=True)
                else:
                    display_checklist_result(invalid_items, "❌ 废标项检测结果", "warning-box")
            
            # 商务得分检查结果
            if result.get("commercial_score_check"):
                display_checklist_result(result["commercial_score_check"], "💰 商务得分检查")
            
            # 技术方案评估结果
            if result.get("technical_plan_check"):
                display_checklist_result(result["technical_plan_check"], "🔧 技术方案评估")
            
            # 指标应答验证结果
            if result.get("indicator_response_check"):
                display_checklist_result(result["indicator_response_check"], "📊 指标应答验证")
            
            # 技术得分点分析结果
            if result.get("technical_score_check"):
                display_checklist_result(result["technical_score_check"], "🎯 技术得分点分析")
            
            # 文件结构检查结果
            if result.get("bid_structure_check"):
                display_checklist_result(result["bid_structure_check"], "📁 文件结构检查")
            
            # 修改建议汇总
            st.markdown('<h2 class="section-header">💡 修改建议汇总</h2>', unsafe_allow_html=True)
            if result.get("modification_summary"):
                summary = result["modification_summary"]
                display_checklist_result(summary, "💡 修改建议汇总", "info-box")
            
            # 下载结果按钮
            st.markdown("---")

            # 下载选项
            col1, col2 = st.columns(2)

            with col1:
                if st.button("📄 下载Word报告"):
                    docx_data = generate_docx_report(result)
                    st.download_button(
                        label="下载DOCX报告",
                        data=docx_data,
                        file_name=f"投标文件分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            with col2:
                if st.button("📥 下载JSON报告"):
                    report_data = json.dumps(result, ensure_ascii=False, indent=2)
                    st.download_button(
                        label="下载JSON报告",
                        data=report_data,
                        file_name="招标文件分析报告.json",
                        mime="application/json"
                    )
            
        except Exception as e:
            st.error(f"分析过程出错: {str(e)}")
            st.error(f"错误详情: {type(e).__name__}")
            import traceback
            st.error(traceback.format_exc())
    
    # 页脚
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 1rem;">
        <p>🤖 安天投标文件智能分析系统 | 基于LangGraph工作流引擎</p>
        <p>💡 AI应用创新激励计划参赛作品</p>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
