"""
生成测试用的招标文件和投标文件
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 创建assets目录（如果不存在）
os.makedirs('assets', exist_ok=True)

# ==================== 生成招标文件 ====================
tender_doc = Document()

# 标题
title = tender_doc.add_heading('网络安全系统采购项目招标文件', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 第一章：项目概况
tender_doc.add_heading('第一章 项目概况', 1)
tender_doc.add_paragraph('项目名称：网络安全系统采购项目')
tender_doc.add_paragraph('项目编号：TC-2024-001')
tender_doc.add_paragraph('预算金额：500万元')

# 第二章：废标项
tender_doc.add_heading('第二章 废标项', 1)
tender_doc.add_paragraph('出现以下情况之一的，视为废标：')
tender_doc.add_paragraph('1. 投标人未提供有效的营业执照', style='List Bullet')
tender_doc.add_paragraph('2. 投标人未提供ISO27001信息安全管理体系认证证书', style='List Bullet')
tender_doc.add_paragraph('3. 投标报价超过预算金额', style='List Bullet')
tender_doc.add_paragraph('4. 投标文件未加盖公章', style='List Bullet')
tender_doc.add_paragraph('5. 技术方案中未提供系统架构图', style='List Bullet')

# 第三章：商务要求
tender_doc.add_heading('第三章 商务要求', 1)
tender_doc.add_paragraph('1. 投标人资质要求：')
tender_doc.add_paragraph('   - 具有独立法人资格', style='List Number')
tender_doc.add_paragraph('   - 注册资金不低于1000万元', style='List Number')
tender_doc.add_paragraph('   - 具有ISO27001认证', style='List Number')
tender_doc.add_paragraph('   - 具有CMMI3级以上认证', style='List Number')

tender_doc.add_paragraph('2. 业绩要求：')
tender_doc.add_paragraph('   - 近三年至少完成3个类似网络安全项目', style='List Number')
tender_doc.add_paragraph('   - 单个项目合同金额不低于200万元', style='List Number')

# 第四章：商务评分规则（30分）
tender_doc.add_heading('第四章 商务评分规则', 1)
tender_doc.add_paragraph('商务部分总分30分，评分标准如下：')
tender_doc.add_paragraph('1. 企业资质（10分）：')
tender_doc.add_paragraph('   - ISO27001认证：5分（提供证书复印件）', style='List Bullet')
tender_doc.add_paragraph('   - CMMI3级认证：3分，CMMI5级：5分', style='List Bullet')
tender_doc.add_paragraph('   - CCRC信息安全服务资质：2分', style='List Bullet')

tender_doc.add_paragraph('2. 业绩情况（10分）：')
tender_doc.add_paragraph('   - 近三年类似网络安全项目业绩：每个项目3分，满分10分', style='List Bullet')
tender_doc.add_paragraph('   - 需提供合同复印件及验收证明', style='List Bullet')

tender_doc.add_paragraph('3. 团队实力（5分）：')
tender_doc.add_paragraph('   - 项目经理具有PMP认证：2分', style='List Bullet')
tender_doc.add_paragraph('   - 技术负责人具有CISSP认证：2分', style='List Bullet')
tender_doc.add_paragraph('   - 核心团队成员不少于10人：1分', style='List Bullet')

tender_doc.add_paragraph('4. 服务承诺（5分）：')
tender_doc.add_paragraph('   - 提供7x24小时服务：2分', style='List Bullet')
tender_doc.add_paragraph('   - 质保期3年以上：3分', style='List Bullet')

# 第五章：技术评分细则（70分）
tender_doc.add_heading('第五章 技术评分细则', 1)
tender_doc.add_paragraph('技术部分总分70分，评分标准如下：')

tender_doc.add_paragraph('1. 技术方案（30分）：')
tender_doc.add_paragraph('   - 方案完整性：10分（包含总体方案、技术架构、实施方案）', style='List Bullet')
tender_doc.add_paragraph('   - 技术先进性：10分（采用主流先进技术，具有前瞻性）', style='List Bullet')
tender_doc.add_paragraph('   - 可行性分析：10分（实施方案详细可行，风险控制到位）', style='List Bullet')

tender_doc.add_paragraph('2. 产品性能（20分）：')
tender_doc.add_paragraph('   - 防火墙性能：支持万兆吞吐量（5分）', style='List Bullet')
tender_doc.add_paragraph('   - 入侵检测：误报率低于1%（5分）', style='List Bullet')
tender_doc.add_paragraph('   - 漏洞扫描：支持2000+漏洞库（5分）', style='List Bullet')
tender_doc.add_paragraph('   - 日志审计：支持百万级日志处理（5分）', style='List Bullet')

tender_doc.add_paragraph('3. 技术指标响应（20分）：')
tender_doc.add_paragraph('   - 必须满足所有技术指标要求', style='List Bullet')
tender_doc.add_paragraph('   - 每一项技术指标需提供详细应答', style='List Bullet')
tender_doc.add_paragraph('   - 需提供产品检测报告或证明材料', style='List Bullet')

# 保存招标文件
tender_doc_path = os.path.join('assets', 'tender_document.docx')
tender_doc.save(tender_doc_path)
print(f"招标文件已生成: {tender_doc_path}")

# ==================== 生成投标文件 ====================
bid_doc = Document()

# 标题
bid_title = bid_doc.add_heading('网络安全系统采购项目投标文件', 0)
bid_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 第一章：投标函
bid_doc.add_heading('第一章 投标函', 1)
bid_doc.add_paragraph('致：招标单位')
bid_doc.add_paragraph('我方愿意参加贵单位网络安全系统采购项目的投标，承诺按照招标文件要求提供产品和服务。')
bid_doc.add_paragraph('投标报价：480万元')
bid_doc.add_paragraph('工期：6个月')

# 第二章：资质证明
bid_doc.add_heading('第二章 资质证明', 1)
bid_doc.add_paragraph('1. 企业法人营业执照（已附复印件）')
bid_doc.add_paragraph('2. ISO27001信息安全管理体系认证证书（已附复印件）')
bid_doc.add_paragraph('3. CMMI3级认证证书（已附复印件）')
bid_doc.add_paragraph('注意：未提供CCRC信息安全服务资质')

# 第三章：业绩介绍
bid_doc.add_heading('第三章 业绩介绍', 1)
bid_doc.add_paragraph('近三年完成2个类似项目：')
bid_doc.add_paragraph('1. 某市政务云安全项目，合同金额180万元')
bid_doc.add_paragraph('2. 某企业网络安全建设项目，合同金额150万元')
bid_doc.add_paragraph('注意：只有2个项目，少于要求的3个项目')

# 第四章：技术方案
bid_doc.add_heading('第四章 技术方案', 1)
bid_doc.add_paragraph('4.1 总体方案')
bid_doc.add_paragraph('本项目采用"纵深防御"的网络安全架构，构建多层次安全防护体系。')

bid_doc.add_paragraph('4.2 技术架构')
bid_doc.add_paragraph('方案包括以下子系统：')
bid_doc.add_paragraph('- 防火墙系统', style='List Bullet')
bid_doc.add_paragraph('- 入侵检测系统', style='List Bullet')
bid_doc.add_paragraph('- 漏洞扫描系统', style='List Bullet')
bid_doc.add_paragraph('- 日志审计系统', style='List Bullet')

bid_doc.add_paragraph('注意：技术方案中缺少详细的系统架构图')

bid_doc.add_paragraph('4.3 实施方案')
bid_doc.add_paragraph('项目实施分为四个阶段：需求调研、方案设计、系统部署、验收交付。')

# 第五章：产品参数
bid_doc.add_heading('第五章 产品参数', 1)
bid_doc.add_paragraph('产品性能指标：')
bid_doc.add_paragraph('- 防火墙：支持万兆吞吐量', style='List Bullet')
bid_doc.add_paragraph('- 入侵检测：误报率约2%（略高于要求的1%）', style='List Bullet')
bid_doc.add_paragraph('- 漏洞扫描：支持1500+漏洞库（少于要求的2000+）', style='List Bullet')
bid_doc.add_paragraph('- 日志审计：支持百万级日志处理', style='List Bullet')

# 第六章：团队介绍
bid_doc.add_heading('第六章 团队介绍', 1)
bid_doc.add_paragraph('- 项目经理：具有5年项目管理经验，PMP认证')
bid_doc.add_paragraph('- 技术负责人：具有10年网络安全经验')
bid_doc.add_paragraph('注意：技术负责人未提供CISSP认证')
bid_doc.add_paragraph('- 核心团队：8人（少于要求的10人）')

# 第七章：服务承诺
bid_doc.add_heading('第七章 服务承诺', 1)
bid_doc.add_paragraph('- 提供5x8小时服务（少于要求的7x24小时）')
bid_doc.add_paragraph('- 质保期：2年（少于要求的3年）')

# 保存投标文件
bid_doc_path = os.path.join('assets', 'bid_document.docx')
bid_doc.save(bid_doc_path)
print(f"投标文件已生成: {bid_doc_path}")

print("\n测试文件生成完成！")
